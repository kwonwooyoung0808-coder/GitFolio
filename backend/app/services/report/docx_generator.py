from docx import Document


def generate_docx(report: dict, output_path) -> str:
    document = Document()
    document.add_heading(report.get("project_name", "GitFolio Report"), level=1)
    document.add_paragraph(report.get("repo", {}).get("full_name", ""))
    document.add_paragraph(report.get("repo", {}).get("description", ""))

    sections = [
        ("주요 업무", report.get("main_task", "")),
        ("담당 역할", report.get("role", "")),
        ("기술 스택", ", ".join(report.get("tech_stack", []))),
        ("업무 기간", report.get("period", "")),
        ("개발 인원", report.get("scale", "")),
        ("상세 내용", report.get("details", report.get("outcome", ""))),
    ]
    for title, value in sections:
        document.add_heading(title, level=2)
        document.add_paragraph(value)

    document.save(str(output_path))
    return str(output_path)
